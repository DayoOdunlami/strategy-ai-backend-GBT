# document_processor.py - Enhanced Document Processing (like your Streamlit version)
import asyncio
import aiofiles
from pathlib import Path
import PyPDF2
import docx
import csv
import io
from typing import List, Dict, Optional
import logging
from config import settings

logger = logging.getLogger(__name__)

class EnhancedDocumentProcessor:
    """
    Enhanced document processor that replicates your Streamlit document_processor.py
    Supports PDF, DOCX, TXT, CSV, and MD files with intelligent chunking
    """
    
    def __init__(self):
        self.chunk_size = settings.DEFAULT_CHUNK_SIZE
        self.chunk_overlap = settings.DEFAULT_CHUNK_OVERLAP
        self.max_chunks = settings.MAX_CHUNKS_PER_DOCUMENT

    async def process_file(
        self, 
        file_path: str, 
        chunk_size: int = None, 
        chunk_overlap: int = None
    ) -> List[Dict]:
        """
        Process uploaded file and return chunks with metadata
        Main entry point like your Streamlit process_file function
        """
        try:
            # Use provided parameters or defaults
            self.chunk_size = chunk_size or self.chunk_size
            self.chunk_overlap = chunk_overlap or self.chunk_overlap
            
            file_ext = Path(file_path).suffix.lower()
            logger.info(f"Processing {file_ext} file: {file_path}")
            
            # Extract text based on file type
            if file_ext == ".pdf":
                text = await self._process_pdf(file_path)
            elif file_ext == ".docx":
                text = await self._process_docx(file_path)
            elif file_ext in [".txt", ".md"]:
                text = await self._process_text(file_path)
            elif file_ext == ".csv":
                text = await self._process_csv(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            if not text.strip():
                raise ValueError("No text content could be extracted from the file")
            
            # Chunk the text
            chunks = await self._chunk_text(text, Path(file_path).name)
            
            logger.info(f"Successfully processed file into {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            raise

    async def process_text(
        self, 
        text: str, 
        source_url: Optional[str] = None,
        title: Optional[str] = None
    ) -> List[Dict]:
        """
        Process raw text (for web scraping)
        Used by the web scraper to process scraped content
        """
        try:
            if not text.strip():
                raise ValueError("Empty text provided")
                
            chunks = await self._chunk_text(text, source_url or "web_content")
            
            # Add source URL to each chunk if provided
            if source_url:
                for chunk in chunks:
                    chunk["source_url"] = source_url
                    chunk["source_type"] = "url"
            
            if title:
                for chunk in chunks:
                    chunk["suggested_title"] = title
                    
            logger.info(f"Processed text into {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing text: {e}")
            raise

    # ============================================================================
    # FILE TYPE PROCESSORS (Private Methods)
    # ============================================================================

    async def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text += f"\n--- Page {page_num + 1} ---\n"
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
                        continue
            
            if not text.strip():
                raise ValueError("No text could be extracted from PDF")
                
            logger.info(f"Extracted {len(text)} characters from PDF")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise ValueError(f"Failed to process PDF file: {str(e)}")

    async def _process_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            if not text.strip():
                raise ValueError("No text could be extracted from DOCX")
                
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise ValueError(f"Failed to process DOCX file: {str(e)}")

    async def _process_text(self, file_path: str) -> str:
        """Read plain text files (TXT, MD)"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                text = await file.read()
                
            if not text.strip():
                raise ValueError("Text file is empty")
                
            logger.info(f"Read {len(text)} characters from text file")
            return text.strip()
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
                    text = await file.read()
                logger.info(f"Read {len(text)} characters from text file (latin-1 encoding)")
                return text.strip()
            except Exception as e:
                logger.error(f"Error reading text file with fallback encoding: {e}")
                raise ValueError("Could not read text file with any encoding")
                
        except Exception as e:
            logger.error(f"Error processing text file: {e}")
            raise ValueError(f"Failed to process text file: {str(e)}")

    async def _process_csv(self, file_path: str) -> str:
        """Convert CSV to readable text format"""
        try:
            text = ""
            with open(file_path, 'r', encoding='utf-8') as file:
                # Detect CSV dialect
                sample = file.read(1024)
                file.seek(0)
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
                
                csv_reader = csv.reader(file, delimiter=delimiter)
                headers = next(csv_reader, [])
                
                if headers:
                    text += "CSV Data Analysis:\n"
                    text += f"Columns: {', '.join(headers)}\n\n"
                    
                    # Read rows (limit for readability)
                    rows = list(csv_reader)
                    total_rows = len(rows)
                    
                    text += f"Total rows: {total_rows}\n\n"
                    text += "Sample data:\n"
                    
                    # Include first 20 rows for analysis
                    for i, row in enumerate(rows[:20]):
                        if len(row) == len(headers):
                            row_data = []
                            for header, value in zip(headers, row):
                                if value.strip():
                                    row_data.append(f"{header}: {value}")
                            if row_data:
                                text += f"Row {i+1}: {'; '.join(row_data)}\n"
                    
                    if total_rows > 20:
                        text += f"\n... and {total_rows - 20} more rows\n"
                    
                    # Add summary statistics
                    text += f"\nData Summary:\n"
                    text += f"- Total columns: {len(headers)}\n"
                    text += f"- Total rows: {total_rows}\n"
                    text += f"- Column names: {', '.join(headers)}\n"
                else:
                    text = "Empty CSV file with no headers"
            
            if not text.strip():
                raise ValueError("CSV file could not be processed")
                
            logger.info(f"Processed CSV into {len(text)} characters of structured text")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error processing CSV: {e}")
            raise ValueError(f"Failed to process CSV file: {str(e)}")

    # ============================================================================
    # TEXT CHUNKING (Intelligent Splitting)
    # ============================================================================

    async def _chunk_text(self, text: str, source_name: str) -> List[Dict]:
        """
        Split text into intelligent, overlapping chunks
        Replicates your Streamlit chunking logic with improvements
        """
        try:
            if len(text) <= self.chunk_size:
                return [{
                    "text": text,
                    "chunk_index": 0,
                    "source_name": source_name,
                    "token_count": self._estimate_tokens(text)
                }]
            
            chunks = []
            start = 0
            chunk_index = 0
            
            # Define sentence boundaries (in order of preference)
            sentence_endings = ['. ', '! ', '? ', '.\n', '!\n', '?\n']
            paragraph_endings = ['\n\n', '\n']
            
            while start < len(text) and chunk_index < self.max_chunks:
                end = start + self.chunk_size
                
                # If we're at the end of the text
                if end >= len(text):
                    chunk_text = text[start:].strip()
                    if chunk_text:
                        chunks.append({
                            "text": chunk_text,
                            "chunk_index": chunk_index,
                            "source_name": source_name,
                            "token_count": self._estimate_tokens(chunk_text),
                            "position": {
                                "start": start,
                                "end": len(text)
                            }
                        })
                    break
                
                # Try to break at sentence boundaries
                best_break = end
                
                # Look for sentence endings
                for ending in sentence_endings:
                    pos = text.rfind(ending, start + self.chunk_size // 2, end)
                    if pos > start:
                        best_break = pos + len(ending)
                        break
                
                # If no sentence ending found, try paragraph breaks
                if best_break == end:
                    for ending in paragraph_endings:
                        pos = text.rfind(ending, start + self.chunk_size // 2, end)
                        if pos > start:
                            best_break = pos + len(ending)
                            break
                
                # If still no good break point, break at word boundary
                if best_break == end:
                    space_pos = text.rfind(' ', start + self.chunk_size // 2, end)
                    if space_pos > start:
                        best_break = space_pos + 1
                
                chunk_text = text[start:best_break].strip()
                if chunk_text:
                    chunks.append({
                        "text": chunk_text,
                        "chunk_index": chunk_index,
                        "source_name": source_name,
                        "token_count": self._estimate_tokens(chunk_text),
                        "position": {
                            "start": start,
                            "end": best_break
                        }
                    })
                    chunk_index += 1
                
                # Move start position with overlap
                start = best_break - self.chunk_overlap
                if start <= chunks[-1]["position"]["start"] if chunks else 0:
                    start = best_break  # Avoid infinite loop
            
            # Log chunking statistics
            total_tokens = sum(chunk["token_count"] for chunk in chunks)
            avg_chunk_size = len(chunks) and total_tokens // len(chunks) or 0
            
            logger.info(f"Created {len(chunks)} chunks, avg size: {avg_chunk_size} tokens")
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error chunking text: {e}")
            raise

    def _estimate_tokens(self, text: str) -> int:
        """
        Estimate token count for text
        Rough approximation: 1 token ≈ 0.75 words
        """
        word_count = len(text.split())
        return int(word_count * 1.3)  # Conservative estimate

    # ============================================================================
    # UTILITY METHODS
    # ============================================================================

    async def validate_file(self, file_path: str) -> Dict[str, any]:
        """Validate file before processing"""
        try:
            path = Path(file_path)
            
            # Check if file exists
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Check file size
            file_size = path.stat().st_size
            max_size_bytes = settings.MAX_FILE_SIZE_MB * 1024 * 1024
            
            if file_size > max_size_bytes:
                raise ValueError(f"File too large: {file_size} bytes (max: {max_size_bytes})")
            
            # Check file extension
            file_ext = path.suffix.lower()
            if file_ext not in settings.ALLOWED_FILE_TYPES:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            return {
                "valid": True,
                "file_size": file_size,
                "file_extension": file_ext,
                "filename": path.name
            }
            
        except Exception as e:
            logger.error(f"File validation failed: {e}")
            return {
                "valid": False,
                "error": str(e)
            }

    async def get_file_info(self, file_path: str) -> Dict[str, any]:
        """Get detailed file information"""
        try:
            path = Path(file_path)
            validation = await self.validate_file(file_path)
            
            if not validation["valid"]:
                return validation
            
            info = {
                **validation,
                "estimated_processing_time": self._estimate_processing_time(
                    validation["file_size"], 
                    validation["file_extension"]
                ),
                "estimated_chunks": self._estimate_chunk_count(validation["file_size"])
            }
            
            return info
            
        except Exception as e:
            logger.error(f"Error getting file info: {e}")
            return {"valid": False, "error": str(e)}

    def _estimate_processing_time(self, file_size: int, file_ext: str) -> int:
        """Estimate processing time in seconds"""
        # Base time by file type
        base_times = {
            ".pdf": 2,
            ".docx": 1,
            ".txt": 0.5,
            ".md": 0.5,
            ".csv": 1
        }
        
        base_time = base_times.get(file_ext, 2)
        
        # Add time based on file size (MB)
        size_mb = file_size / (1024 * 1024)
        size_time = size_mb * 0.5  # 0.5 seconds per MB
        
        return int(base_time + size_time)

    def _estimate_chunk_count(self, file_size: int) -> int:
        """Estimate number of chunks that will be created"""
        # Rough estimate: 1KB ≈ 1 chunk (depending on text density)
        size_kb = file_size / 1024
        estimated_chunks = max(1, int(size_kb / (self.chunk_size / 1000)))
        return min(estimated_chunks, self.max_chunks)