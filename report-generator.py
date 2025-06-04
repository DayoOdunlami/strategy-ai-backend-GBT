# report_generator.py - Comprehensive Report Generation System
import asyncio
from datetime import datetime
from typing import Dict, List, Optional, Any, Union
import logging
import io
import os
import uuid
from pathlib import Path
import json

# Document generation libraries
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import settings
from specialized_agents import OrchestrationAgent

logger = logging.getLogger(__name__)

class ReportGenerator:
    """
    Comprehensive report generation system
    Creates PDF and DOCX reports with professional formatting
    """
    
    def __init__(self):
        self.orchestrator = OrchestrationAgent()
        self.reports_dir = Path("reports")
        self.reports_dir.mkdir(exist_ok=True)
        
        # Report templates
        self.report_templates = {
            "strategy_analysis": {
                "title": "Strategic Analysis Report",
                "sections": [
                    "Executive Summary",
                    "Current Situation Analysis", 
                    "Strategic Assessment",
                    "Key Findings",
                    "Recommendations",
                    "Implementation Roadmap",
                    "Risk Assessment",
                    "Success Metrics"
                ]
            },
            "project_similarity": {
                "title": "Project Similarity Analysis",
                "sections": [
                    "Analysis Overview",
                    "Similar Projects Identified",
                    "Comparative Analysis",
                    "Lessons Learned",
                    "Success Factors",
                    "Risk Patterns",
                    "Recommendations"
                ]
            },
            "lessons_learned": {
                "title": "Lessons Learned Report",
                "sections": [
                    "Executive Summary",
                    "Methodology",
                    "Key Lessons by Category",
                    "Success Stories",
                    "Failure Analysis",
                    "Best Practices",
                    "Implementation Guidelines"
                ]
            },
            "sector_performance": {
                "title": "Sector Performance Analysis",
                "sections": [
                    "Executive Summary",
                    "Performance Metrics",
                    "Comparative Analysis",
                    "Trend Analysis",
                    "Key Performance Indicators",
                    "Recommendations",
                    "Action Plan"
                ]
            }
        }
        
        logger.info("Initialized report generator with AI orchestration")

    async def generate_report(
        self,
        report_type: str,
        parameters: Dict[str, Any],
        format: str = "pdf"  # pdf, docx, both
    ) -> Dict[str, Any]:
        """
        Main report generation function
        Creates comprehensive reports using AI agents
        """
        try:
            report_id = str(uuid.uuid4())
            logger.info(f"Generating {report_type} report in {format} format")
            
            # Step 1: Gather data and context
            context_data = await self._gather_report_context(report_type, parameters)
            
            # Step 2: Generate report content using AI agents
            report_content = await self._generate_report_content(
                report_type, context_data, parameters
            )
            
            # Step 3: Create formatted documents
            generated_files = []
            
            if format in ["pdf", "both"]:
                pdf_path = await self._generate_pdf_report(
                    report_id, report_type, report_content, parameters
                )
                if pdf_path:
                    generated_files.append({
                        "format": "pdf",
                        "path": pdf_path,
                        "filename": f"{report_type}_{report_id}.pdf"
                    })
            
            if format in ["docx", "both"]:
                docx_path = await self._generate_docx_report(
                    report_id, report_type, report_content, parameters
                )
                if docx_path:
                    generated_files.append({
                        "format": "docx", 
                        "path": docx_path,
                        "filename": f"{report_type}_{report_id}.docx"
                    })
            
            # Step 4: Generate metadata
            report_metadata = {
                "report_id": report_id,
                "report_type": report_type,
                "generated_at": datetime.now().isoformat(),
                "parameters": parameters,
                "files": generated_files,
                "word_count": len(report_content.get("full_content", "").split()),
                "sections_count": len(report_content.get("sections", {})),
                "ai_agents_used": report_content.get("agents_used", [])
            }
            
            # Save metadata
            await self._save_report_metadata(report_id, report_metadata)
            
            logger.info(f"Successfully generated report {report_id}")
            return {
                "success": True,
                "report_id": report_id,
                "metadata": report_metadata,
                "download_urls": [f"/api/reports/{report_id}/download/{file['filename']}" for file in generated_files]
            }
            
        except Exception as e:
            logger.error(f"Error generating report: {e}")
            return {
                "success": False,
                "error": str(e),
                "report_id": None
            }

    async def _gather_report_context(self, report_type: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """Gather context data for report generation"""
        
        context = {
            "report_type": report_type,
            "sector": parameters.get("sector", "General"),
            "date_range": parameters.get("date_range", {}),
            "scope": parameters.get("scope", "comprehensive"),
            "documents": [],
            "analytics": {}
        }
        
        # Add specific context based on report type
        if report_type == "strategy_analysis":
            context.update({
                "focus_areas": parameters.get("focus_areas", []),
                "strategic_objectives": parameters.get("objectives", []),
                "stakeholders": parameters.get("stakeholders", [])
            })
        
        elif report_type == "project_similarity":
            context.update({
                "target_project": parameters.get("target_project", {}),
                "similarity_criteria": parameters.get("criteria", []),
                "comparison_metrics": parameters.get("metrics", [])
            })
        
        elif report_type == "lessons_learned":
            context.update({
                "project_scope": parameters.get("project_scope", ""),
                "time_period": parameters.get("time_period", ""),
                "categories": parameters.get("categories", [])
            })
        
        return context

    async def _generate_report_content(
        self, 
        report_type: str, 
        context_data: Dict[str, Any], 
        parameters: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate report content using specialized AI agents"""
        
        # Prepare request for AI agents
        ai_request = {
            "type": "report",
            "report_type": report_type,
            "context": json.dumps(context_data),
            "parameters": parameters,
            "sector": context_data.get("sector", "General")
        }
        
        # Generate content using orchestration agent
        ai_response = await self.orchestrator.process_request(ai_request)
        
        # Structure the content
        if "structured_report" in ai_response:
            sections = ai_response["structured_report"].get("sections", {})
        else:
            # Parse unstructured response into sections
            sections = self._parse_response_into_sections(
                ai_response.get("response", ""), 
                report_type
            )
        
        return {
            "full_content": ai_response.get("response", ""),
            "sections": sections,
            "metadata": {
                "generated_by": ai_response.get("agent", "ReportAgent"),
                "confidence": ai_response.get("confidence", 0.0),
                "agents_used": ai_response.get("agents_used", [])
            },
            "template": self.report_templates.get(report_type, {})
        }

    def _parse_response_into_sections(self, content: str, report_type: str) -> Dict[str, str]:
        """Parse AI response into structured sections"""
        
        sections = {}
        current_section = "introduction"
        current_content = []
        
        template = self.report_templates.get(report_type, {})
        expected_sections = template.get("sections", [])
        
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # Check if this line is a section header
            is_header = False
            for expected_section in expected_sections:
    def _parse_response_into_sections(self, content: str, report_type: str) -> Dict[str, str]:
        """Parse AI response into structured sections"""
        
        sections = {}
        current_section = "introduction"
        current_content = []
        
        template = self.report_templates.get(report_type, {})
        expected_sections = template.get("sections", [])
        
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # Check if this line is a section header
            is_header = False
            for expected_section in expected_sections:
                if expected_section.upper() in line.upper() and len(line) < 100:
                    # Save previous section
                    if current_content:
                        sections[current_section] = '\n'.join(current_content)
                    
                    # Start new section
                    current_section = expected_section.lower().replace(' ', '_')
                    current_content = []
                    is_header = True
                    break
            
            if not is_header:
                current_content.append(line)
        
        # Save final section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections

    async def _generate_pdf_report(
        self, 
        report_id: str, 
        report_type: str, 
        content: Dict[str, Any], 
        parameters: Dict[str, Any]
    ) -> Optional[str]:
        """Generate PDF report with professional formatting"""
        
        try:
            filename = f"{report_type}_{report_id}.pdf"
            filepath = self.reports_dir / filename
            
            # Create PDF document
            doc = SimpleDocTemplate(
                str(filepath),
                pagesize=A4,
                rightMargin=inch,
                leftMargin=inch,
                topMargin=inch,
                bottomMargin=inch
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=18,
                spaceAfter=30,
                alignment=TA_CENTER,
                textColor=colors.HexColor('#006E51')
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading1'],
                fontSize=14,
                spaceAfter=12,
                spaceBefore=20,
                textColor=colors.HexColor('#006E51')
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                alignment=TA_JUSTIFY
            )
            
            # Build document content
            story = []
            
            # Title page
            template = content.get("template", {})
            report_title = template.get("title", report_type.replace('_', ' ').title())
            story.append(Paragraph(report_title, title_style))
            story.append(Spacer(1, 20))
            
            # Metadata table
            metadata_data = [
                ["Report Type", report_type.replace('_', ' ').title()],
                ["Generated", datetime.now().strftime("%B %d, %Y")],
                ["Sector", parameters.get("sector", "General")],
                ["Report ID", report_id]
            ]
            
            metadata_table = Table(metadata_data, colWidths=[2*inch, 3*inch])
            metadata_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F5F5F5')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(metadata_table)
            story.append(PageBreak())
            
            # Table of contents
            story.append(Paragraph("Table of Contents", heading_style))
            sections = content.get("sections", {})
            
            for i, section_name in enumerate(sections.keys(), 1):
                section_title = section_name.replace('_', ' ').title()
                story.append(Paragraph(f"{i}. {section_title}", body_style))
            
            story.append(PageBreak())
            
            # Report sections
            for section_name, section_content in sections.items():
                if section_content.strip():
                    section_title = section_name.replace('_', ' ').title()
                    story.append(Paragraph(section_title, heading_style))
                    
                    # Split content into paragraphs
                    paragraphs = section_content.split('\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            story.append(Paragraph(paragraph.strip(), body_style))
                    
                    story.append(Spacer(1, 20))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"Generated PDF report: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generating PDF report: {e}")
            return None

    async def _generate_docx_report(
        self, 
        report_id: str, 
        report_type: str, 
        content: Dict[str, Any], 
        parameters: Dict[str, Any]
    ) -> Optional[str]:
        """Generate DOCX report with professional formatting"""
        
        try:
            filename = f"{report_type}_{report_id}.docx"
            filepath = self.reports_dir / filename
            
            # Create document
            doc = docx.Document()
            
            # Set document styles
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = docx.shared.Pt(11)
            
            # Title
            template = content.get("template", {})
            report_title = template.get("title", report_type.replace('_', ' ').title())
            title = doc.add_heading(report_title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_heading('Report Information', level=1)
            
            metadata_table = doc.add_table(rows=4, cols=2)
            metadata_table.style = 'Table Grid'
            
            metadata_data = [
                ("Report Type", report_type.replace('_', ' ').title()),
                ("Generated", datetime.now().strftime("%B %d, %Y")),
                ("Sector", parameters.get("sector", "General")),
                ("Report ID", report_id)
            ]
            
            for i, (key, value) in enumerate(metadata_data):
                row_cells = metadata_table.rows[i].cells
                row_cells[0].text = key
                row_cells[1].text = str(value)
            
            doc.add_page_break()
            
            # Table of contents
            doc.add_heading('Table of Contents', level=1)
            sections = content.get("sections", {})
            
            for i, section_name in enumerate(sections.keys(), 1):
                section_title = section_name.replace('_', ' ').title()
                toc_paragraph = doc.add_paragraph(f"{i}. {section_title}")
                toc_paragraph.style = 'List Number'
            
            doc.add_page_break()
            
            # Report sections
            for section_name, section_content in sections.items():
                if section_content.strip():
                    section_title = section_name.replace('_', ' ').title()
                    doc.add_heading(section_title, level=1)
                    
                    # Add content paragraphs
                    paragraphs = section_content.split('\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            doc.add_paragraph(paragraph.strip())
            
            # Save document
            doc.save(str(filepath))
            
            logger.info(f"Generated DOCX report: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generating DOCX report: {e}")
            return None

    async def _save_report_metadata(self, report_id: str, metadata: Dict[str, Any]):
        """Save report metadata to file"""
        try:
            metadata_file = self.reports_dir / f"{report_id}_metadata.json"
            with open(metadata_file, 'w') as f:
                json.dump(metadata, f, indent=2)
            
            logger.info(f"Saved report metadata: {metadata_file}")
            
        except Exception as e:
            logger.error(f"Error saving report metadata: {e}")

    async def get_report_file(self, report_id: str, filename: str) -> Optional[str]:
        """Get report file path for download"""
        try:
            filepath = self.reports_dir / filename
            if filepath.exists():
                return str(filepath)
            return None
            
        except Exception as e:
            logger.error(f"Error getting report file: {e}")
            return None

    async def get_report_metadata(self, report_id: str) -> Optional[Dict[str, Any]]:
        """Get report metadata"""
        try:
            metadata_file = self.reports_dir / f"{report_id}_metadata.json"
            if metadata_file.exists():
                with open(metadata_file, 'r') as f:
                    return json.load(f)
            return None
            
        except Exception as e:
            logger.error(f"Error getting report metadata: {e}")
            return None

    async def list_reports(self, limit: int = 50) -> List[Dict[str, Any]]:
        """List all generated reports"""
        try:
            reports = []
            
            # Find all metadata files
            for metadata_file in self.reports_dir.glob("*_metadata.json"):
                try:
                    with open(metadata_file, 'r') as f:
                        metadata = json.load(f)
                        
                    # Add file existence check
                    files_exist = []
                    for file_info in metadata.get("files", []):
                        filepath = Path(file_info["path"])
                        files_exist.append({
                            **file_info,
                            "exists": filepath.exists(),
                            "size_mb": round(filepath.stat().st_size / (1024*1024), 2) if filepath.exists() else 0
                        })
                    
                    metadata["files"] = files_exist
                    reports.append(metadata)
                    
                except Exception as e:
                    logger.warning(f"Error reading metadata file {metadata_file}: {e}")
                    continue
            
            # Sort by generation date (newest first)
            reports.sort(key=lambda x: x.get("generated_at", ""), reverse=True)
            
            return reports[:limit]
            
        except Exception as e:
            logger.error(f"Error listing reports: {e}")
            return []

    async def delete_report(self, report_id: str) -> bool:
        """Delete a report and its files"""
        try:
            deleted_files = 0
            
            # Delete report files
            for file_path in self.reports_dir.glob(f"*{report_id}*"):
                try:
                    file_path.unlink()
                    deleted_files += 1
                except Exception as e:
                    logger.warning(f"Error deleting file {file_path}: {e}")
            
            logger.info(f"Deleted {deleted_files} files for report {report_id}")
            return deleted_files > 0
            
        except Exception as e:
            logger.error(f"Error deleting report {report_id}: {e}")
            return False

    def get_available_report_types(self) -> List[Dict[str, Any]]:
        """Get information about available report types"""
        return [
            {
                "type": report_type,
                "title": template["title"],
                "sections": template["sections"],
                "description": self._get_report_description(report_type)
            }
            for report_type, template in self.report_templates.items()
        ]

    def _get_report_description(self, report_type: str) -> str:
        """Get description for report type"""
        descriptions = {
            "strategy_analysis": "Comprehensive strategic analysis with recommendations and implementation roadmap",
            "project_similarity": "Comparative analysis identifying similar projects with lessons learned",
            "lessons_learned": "Structured compilation of project insights, successes, and improvement areas",
            "sector_performance": "Performance metrics and trend analysis for specific sectors"
        }
        return descriptions.get(report_type, "Customized report based on your requirements")

class ReportTemplateManager:
    """
    Manages custom report templates
    Allows users to create and modify report structures
    """
    
    def __init__(self):
        self.templates_dir = Path("report_templates")
        self.templates_dir.mkdir(exist_ok=True)
    
    async def create_custom_template(
        self, 
        template_name: str, 
        template_config: Dict[str, Any]
    ) -> bool:
        """Create a custom report template"""
        try:
            template_file = self.templates_dir / f"{template_name}.json"
            
            # Validate template structure
            required_fields = ["title", "sections"]
            if not all(field in template_config for field in required_fields):
                logger.error(f"Template missing required fields: {required_fields}")
                return False
            
            # Save template
            with open(template_file, 'w') as f:
                json.dump(template_config, f, indent=2)
            
            logger.info(f"Created custom template: {template_name}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating custom template: {e}")
            return False
    
    async def get_custom_template(self, template_name: str) -> Optional[Dict[str, Any]]:
        """Get a custom report template"""
        try:
            template_file = self.templates_dir / f"{template_name}.json"
            
            if template_file.exists():
                with open(template_file, 'r') as f:
                    return json.load(f)
            
            return None
            
        except Exception as e:
            logger.error(f"Error getting custom template: {e}")
            return None
    
    async def list_custom_templates(self) -> List[Dict[str, Any]]:
        """List all custom templates"""
        try:
            templates = []
            
            for template_file in self.templates_dir.glob("*.json"):
                try:
                    with open(template_file, 'r') as f:
                        template_data = json.load(f)
                    
                    templates.append({
                        "name": template_file.stem,
                        "title": template_data.get("title", "Custom Template"),
                        "sections": len(template_data.get("sections", [])),
                        "created": datetime.fromtimestamp(template_file.stat().st_mtime).isoformat()
                    })
                    
                except Exception as e:
                    logger.warning(f"Error reading template {template_file}: {e}")
                    continue
            
            return templates
            
        except Exception as e:
            logger.error(f"Error listing custom templates: {e}")
            return []